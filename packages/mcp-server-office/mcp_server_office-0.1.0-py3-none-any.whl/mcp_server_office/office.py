import os
from typing import Dict
from docx import Document
from docx.table import Table
from docx.oxml import OxmlElement
from mcp.server.lowlevel import Server, NotificationOptions
from mcp.server.stdio import stdio_server
from mcp.server.models import InitializationOptions
from mcp import types
import difflib
from mcp_server_office.tools import READ_DOCX, WRITE_DOCX, EDIT_DOCX
import uuid

server = Server("office-server")

async def validate_path(path: str) -> bool:
    if not os.path.isabs(path):
        raise ValueError(f"Not a absolute path: {path}")
    if not os.path.isfile(path):
        raise ValueError(f"File not found: {path}")
    elif path.endswith(".docx"):
        return True
    else:
        return False

def extract_table_text(table: Table) -> str:
    """Extract text from table with formatting."""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(" | ".join(cells))
    return "\n".join(rows)

def process_track_changes(element: OxmlElement) -> str:
    """Process track changes in a paragraph element."""
    text = ""
    for child in element:
        if child.tag.endswith('r'):  # Normal run
            for run_child in child:
                if run_child.tag.endswith('t'):
                    text += run_child.text if run_child.text else ""
        elif child.tag.endswith('ins'):  # Insertion
            inserted_text = ""
            for run in child.findall('.//w:t', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                inserted_text += run.text if run.text else ""
            if inserted_text:
                text += inserted_text
    return text

async def read_docx(path: str) -> str:
    """Read docx file as text including tables.
    
    Args:
        path: relative path to target docx file
    Returns:
        str: Text representation of the document including tables
    """
    if not await validate_path(path):
        raise ValueError(f"Not a docx file: {path}")
    
    document = Document(path)
    content = []

    paragraph_index = 0
    table_index = 0
    
    # Process all elements in order
    for element in document._body._body:
        # Process paragraph
        if element.tag.endswith('p'):
            paragraph = document.paragraphs[paragraph_index]
            paragraph_index += 1
            # Check for image
            if paragraph._element.findall('.//w:drawing', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                content.append("[Image]")
            # Check for text
            else:
                text = process_track_changes(paragraph._element)
                if text.strip():
                    content.append(text)
        # Process table
        elif element.tag.endswith('tbl'):
            table = document.tables[table_index]
            table_index += 1
            table_text = extract_table_text(table)
            content.append(f"[Table]\n{table_text}")

    return "\n\n".join(content)

async def write_docx(path: str, content: str) -> None:
    """Create a new docx file with the given content.
    
    Args:
        path: target path to create docx file
        content: text content to write
    """
    document = Document()
    
    # Split content into sections
    sections = content.split("\n\n")
    
    for section in sections:
        if section.startswith("[Table]"):
            # Create table from text representation
            table_content = section[7:].strip()  # Remove [Table] prefix
            rows = table_content.split("\n")
            if rows:
                num_columns = len(rows[0].split(" | "))
                table = document.add_table(rows=len(rows), cols=num_columns)
                
                for i, row in enumerate(rows):
                    cells = row.split(" | ")
                    for j, cell in enumerate(cells):
                        table.cell(i, j).text = cell.strip()
        elif section.startswith("[Image]"):
            document.add_paragraph("[Image placeholder]")
        else:
            document.add_paragraph(section)
    
    document.save(path)

async def edit_docx(path: str, edits: list[Dict[str, str]]) -> str:
    """Edit docx file by replacing multiple text occurrences.
    
    Args:
        path: path to target docx file
        edits: list of dictionaries containing 'search' and 'replace' pairs
            [{'search': 'text to find', 'replace': 'text to replace with'}, ...]
    Returns:
        str: A git-style diff showing the changes made
    """
    if not await validate_path(path):
        raise ValueError(f"Not a docx file: {path}")
    
    # Read original content with track changes
    original = await read_docx(path)
    modified = str(original)
    not_found = []
    
    markers ={edit["search"]: uuid.uuid4().hex for edit in edits}
    
    for search, id in markers.items():
        if not search in original:
            not_found.append(search)
            continue
       
        modified = modified.replace(search, id)
    
    if not_found:
        raise ValueError(f"Search text not found: {', '.join(not_found)}")
    
    for edit in edits:
        search = edit["search"]
        replace = edit["replace"]
        modified = modified.replace(markers[search], replace)
    
    # Write modifications
    await write_docx(path, modified)
    
    # Read modified content and create diff
    modified = await read_docx(path)
    result = "\n".join([line for line in difflib.unified_diff(original.split("\n"), modified.split("\n"))])
    return result

@server.list_tools()
async def list_tools() -> list[types.Tool]:
    return [READ_DOCX, EDIT_DOCX, WRITE_DOCX]

@server.call_tool()
async def call_tool(
    name: str,
    arguments: dict
) -> list[types.TextContent]:
    if name == "read_docx":
        content = await read_docx(arguments["path"])
        return [types.TextContent(type="text", text=content)]
    elif name == "write_docx":
        await write_docx(arguments["path"], arguments["content"])
        return [types.TextContent(type="text", text="Document created successfully")]
    elif name == "edit_docx":
        result = await edit_docx(arguments["path"], arguments["edits"])
        return [types.TextContent(type="text", text=result)]
    raise ValueError(f"Tool not found: {name}")

async def run():
    async with stdio_server() as (read_stream, write_stream):
        await server.run(
            read_stream,
            write_stream,
            InitializationOptions(
                server_name="office-file-server",
                server_version="0.1.0",
                capabilities=server.get_capabilities(
                    notification_options=NotificationOptions(),
                    experimental_capabilities={},
                ),
            )
        )

if __name__ == "__main__":
    import asyncio
    asyncio.run(run())
