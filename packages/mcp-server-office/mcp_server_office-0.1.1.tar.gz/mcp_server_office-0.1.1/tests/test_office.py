import os
import pytest
from mcp_server_office.office import validate_path, read_docx, write_docx, edit_docx, extract_table_text
from docx import Document
from docx.table import Table
from docx.oxml.shared import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

@pytest.fixture
def sample_docx_with_track_changes():
    """Create a sample docx file with track changes for testing."""
    path = "test_track_changes.docx"
    doc = Document()
    paragraph = doc.add_paragraph()
    
    # Add text with track changes
    run = OxmlElement('w:r')
    text = OxmlElement('w:t')
    text.text = "Original"
    run.append(text)
    paragraph._element.append(run)
    
    # Add deletion
    del_element = OxmlElement('w:del')
    del_element.set(qn('w:author'), 'Test Author')
    del_element.set(qn('w:date'), '2024-01-27T00:00:00Z')
    del_run = OxmlElement('w:r')
    del_text = OxmlElement('w:delText')
    del_text.text = " deleted"
    del_run.append(del_text)
    del_element.append(del_run)
    paragraph._element.append(del_element)
    
    # Add insertion
    ins_element = OxmlElement('w:ins')
    ins_element.set(qn('w:author'), 'Test Author')
    ins_element.set(qn('w:date'), '2024-01-27T00:00:00Z')
    ins_run = OxmlElement('w:r')
    ins_text = OxmlElement('w:t')
    ins_text.text = " inserted"
    ins_run.append(ins_text)
    ins_element.append(ins_run)
    paragraph._element.append(ins_element)
    
    doc.save(path)
    yield path
    if os.path.exists(path):
        os.remove(path)

@pytest.fixture
def sample_docx():
    """Create a sample docx file for testing."""
    path = "test_sample.docx"
    doc = Document()
    doc.add_paragraph("Hello World")
    
    # Add table
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A1"
    table.cell(0, 1).text = "B1"
    table.cell(1, 0).text = "A2"
    table.cell(1, 1).text = "B2"
    
    doc.add_paragraph("Goodbye World")
    doc.save(path)
    yield path
    # Cleanup
    if os.path.exists(path):
        os.remove(path)

@pytest.mark.asyncio
async def test_validate_path(sample_docx):
    """Test path validation."""
    assert await validate_path(os.path.abspath(sample_docx)) == True
    with pytest.raises(ValueError):
        await validate_path(sample_docx) 
    with pytest.raises(ValueError):
        await validate_path("nonexistent.docx")

@pytest.mark.asyncio
async def test_read_docx_with_track_changes(sample_docx_with_track_changes):
    """Test reading docx file with track changes."""
    content = await read_docx(os.path.abspath(sample_docx_with_track_changes))
    assert "Original" in content
    assert not "deleted" in content
    assert "inserted" in content

@pytest.mark.asyncio
async def test_read_docx(sample_docx):
    """Test reading docx file."""
    content = await read_docx(os.path.abspath(sample_docx))
    assert "Hello World" in content
    assert "Goodbye World" in content
    assert "[Table]" in content
    assert "A1 | B1" in content
    assert "A2 | B2" in content

    with pytest.raises(ValueError):
        await read_docx("nonexistent.docx")

@pytest.mark.asyncio
async def test_write_docx():
    """Test writing docx file."""
    test_path = "test_write.docx"
    test_content = "Test Paragraph\n\n[Table]\nC1 | D1\nC2 | D2\n\nFinal Paragraph"
    
    try:
        await write_docx(test_path, test_content)
        assert os.path.exists(test_path)
        
        # Verify content
        doc = Document(test_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text]
        assert "Test Paragraph" in paragraphs
        assert "Final Paragraph" in paragraphs
        
        # Verify table
        table = doc.tables[0]
        assert table.cell(0, 0).text == "C1"
        assert table.cell(0, 1).text == "D1"
        assert table.cell(1, 0).text == "C2"
        assert table.cell(1, 1).text == "D2"
    
    finally:
        if os.path.exists(test_path):
            os.remove(test_path)

@pytest.mark.asyncio
async def test_edit_docx_with_track_changes(sample_docx_with_track_changes):
    """Test editing docx file with track changes."""
    abs_path = os.path.abspath(sample_docx_with_track_changes)
    result = await edit_docx(abs_path, [{"search": "Original", "replace": "Modified"}])
    
    # Verify track changes are preserved
    content = await read_docx(abs_path)
    assert not "Original" in content
    assert "Modified" in content
    assert not "deleted" in content
    assert "inserted" in content

@pytest.fixture
def complex_docx():
    """Create a sample docx file with complex content for testing cross-paragraph and table edits."""
    path = "test_complex.docx"
    doc = Document()
    
    # Add paragraphs with text that spans multiple paragraphs
    doc.add_paragraph("First part of")
    doc.add_paragraph("a sentence that")
    doc.add_paragraph("spans multiple paragraphs")
    
    doc.add_paragraph("Some text before table")
    
    # Add table with text that will be edited
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Table"
    table.cell(0, 1).text = "Content"
    table.cell(1, 0).text = "More"
    table.cell(1, 1).text = "Text"
    
    doc.add_paragraph("Some text after table")
    
    doc.save(path)
    yield path
    if os.path.exists(path):
        os.remove(path)

@pytest.mark.asyncio
async def test_edit_docx(sample_docx):
    abs_sample_docx = os.path.abspath(sample_docx)
    
    """Test editing docx file."""
    # Test single edit
    result = await edit_docx(abs_sample_docx, [{"search": "Hello", "replace": "Hi"}])
    assert "-Hello" in result
    assert "+Hi" in result
    
    # Test multiple edits
    result = await edit_docx(abs_sample_docx, [
        {"search": "Hi", "replace": "Hellow"},
        {"search": "World", "replace": "Everyone"}
    ])
    assert "-Hi World" in result
    assert "+Hellow Everyone" in result
    
    # Test non-existent text
    with pytest.raises(ValueError) as exc_info:
        await edit_docx(abs_sample_docx, [{"search": "NonexistentText", "replace": "NewText"}])
    assert "Search text not found: NonexistentText" in str(exc_info.value)
    
    # Test multiple edits with one non-existent
    with pytest.raises(ValueError) as exc_info:
        await edit_docx(abs_sample_docx, [
            {"search": "Hello", "replace": "Hi"},
            {"search": "NonexistentText", "replace": "NewText"}
        ])
    assert "Search text not found: NonexistentText" in str(exc_info.value)
    
    # Test invalid file
    with pytest.raises(ValueError):
        await edit_docx("nonexistent.docx", [{"search": "Hello", "replace": "Hi"}])

@pytest.mark.asyncio
async def test_edit_docx_cross_paragraph(complex_docx):
    """Test editing text that spans multiple paragraphs."""
    abs_complex_docx = os.path.abspath(complex_docx)
    
    # Test editing text that spans paragraphs
    result = await edit_docx(abs_complex_docx, [
        {"search": "First part of\n\na sentence that\n\nspans multiple paragraphs", "replace": "The beginning of"}
    ])
    assert "-First part of" in result
    assert "-a sentence that" in result
    assert "-spans multiple paragraphs" in result
    assert "+The beginning of" in result
    
    # Verify content
    content = await read_docx(abs_complex_docx)
    assert "The beginning of" in content

@pytest.mark.asyncio
async def test_edit_docx_table_content(complex_docx):
    """Test editing text within table cells."""
    abs_complex_docx = os.path.abspath(complex_docx)
    
    # Test editing table content
    result = await edit_docx(abs_complex_docx, [
        {"search": "Table | Content", "replace": "Modified | Cell"}
    ])
    assert "-Table | Content" in result
    assert "+Modified | Cell" in result
    
    # Test editing text that includes table and regular paragraph
    result = await edit_docx(abs_complex_docx, [
        {"search": "Some text before table\n\n[Table]\nModified | Cell\nMore | Text\n\nSome text after", 
         "replace": "Updated content with table"}
    ])
    assert "-Some text before table" in result
    assert "+Updated content with table" in result

def test_extract_table_text():
    """Test table text extraction."""
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "X1"
    table.cell(0, 1).text = "Y1"
    table.cell(1, 0).text = "X2"
    table.cell(1, 1).text = "Y2"
    
    result = extract_table_text(table)
    assert "X1 | Y1" in result
    assert "X2 | Y2" in result
