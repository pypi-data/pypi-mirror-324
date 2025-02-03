from docx.document import Document
from docx.text.paragraph import Paragraph
from typing import Dict, List, Tuple

def get_symbol_map() -> Dict[str, str]:
  return {
    '0xf020':' ', '0xf021':'!', '0xf022':'∀', '0xf023':'#',
    '0xf024':'∃', '0xf025':'%', '0xf026':'&', '0xf027':'∍', 
    '0xf028':'(', '0xf029':')', '0xf02a':'*', '0xf02b':'+', 
    '0xf02c':',', '0xf02d':'-', '0xf02e':'.', '0xf02f':'/', 
    '0xf030':'0', '0xf031':'1', '0xf032':'2', '0xf033':'3', 
    '0xf034':'4', '0xf035':'5', '0xf036':'6', '0xf037':'7', 
    '0xf038':'8', '0xf039':'9', '0xf03a':':', '0xf03b':';', 
    '0xf03c':'<', '0xf03d':'=', '0xf03e':'>', '0xf03f':'?', 
    '0xf040':'≅', '0xf041':'Α', '0xf042':'Β', '0xf043':'Χ', 
    '0xf044':'Δ', '0xf045':'Ε', '0xf046':'Φ', '0xf047':'Γ', 
    '0xf048':'Η', '0xf049':'Ι', '0xf04a':'ϑ', '0xf04b':'Κ', 
    '0xf04c':'Λ', '0xf04d':'Μ', '0xf04e':'Ν', '0xf04f':'Ο', 
    '0xf050':'Π', '0xf051':'Θ', '0xf052':'Ρ', '0xf053':'Σ', 
    '0xf054':'Τ', '0xf055':'Υ', '0xf056':'ϛ', '0xf057':'Ω', 
    '0xf058':'Ξ', '0xf059':'Ψ', '0xf05a':'Ζ', '0xf05b':'[', 
    '0xf05c':'∴', '0xf05d':']', '0xf05e':'⊥', '0xf05f':'_', 
    '0xf060':'⁻', '0xf061':'α', '0xf062':'β', '0xf063':'χ', 
    '0xf064':'δ', '0xf065':'ε', '0xf066':'ϕ', '0xf067':'γ', 
    '0xf068':'η', '0xf069':'ι', '0xf06a':'φ', '0xf06b':'κ', 
    '0xf06c':'λ', '0xf06d':'μ', '0xf06e':'ν', '0xf06f':'ο', 
    '0xf070':'π', '0xf071':'θ', '0xf072':'ρ', '0xf073':'σ', 
    '0xf074':'τ', '0xf075':'υ', '0xf076':'ϖ', '0xf077':'ω', 
    '0xf078':'ξ', '0xf079':'ψ', '0xf07a':'ζ', '0xf07b':'{', 
    '0xf07c':'|', '0xf07d':'}', '0xf07e':'~', '0xf07f':' ', 
    '0xf0a2':'′', '0xf0a3':'≤', '0xf0a4':'/', '0xf0a5':'∞', 
    '0xf0a6':'f', '0xf0a7':'♣', '0xf0a8':'♦', '0xf0a9':'♥',
    '0xf0aa':'♠', '0xf0ab':'↔', '0xf0ac':'←', '0xf0ad':'↑', 
    '0xf0ae':'→', '0xf0af':'↓',
    '0xf0b0':'°', '0xf0b1':'±', '0xf0b2':'ʺ', '0xf0b3':'≥', 
    '0xf0b4':'×', '0xf0b5':'∝', '0xf0b6':'∂', '0xf0b7':'•', 
    '0xf0b8':'÷', '0xf0b9':'≠', '0xf0ba':'≡', '0xf0bb':'≈', 
    '0xf0bc':'…', '0xf0bd':'|', '0xf0be':'―', '0xf0bf':'↲',
    '0xf0c4':'⊗', '0xf0c5':'⊕', '0xf0c6':'∅', '0xf0c7':'∩', 
    '0xf0c8':'∪', '0xf0c9':'⊃', '0xf0ca':'⊇', '0xf0cb':'⊄', 
    '0xf0cc':'⊂', '0xf0cd':'⊆', '0xf0ce':'∈', '0xf0cf':'∉',
    '0xf0d0':'∠', '0xf0d1':'∇', '0xf0d2':'®', '0xf0d3':'©', 
    '0xf0d4':'™', '0xf0d5':'∏', '0xf0d6':'√', '0xf0d7':'·', 
    '0xf0d8':'¬', '0xf0d9':'∧', '0xf0da':'∨', '0xf0db':'⇔', 
    '0xf0dc':'⇐', '0xf0dd':'⇑', '0xf0de':'⇒', '0xf0df':'⇓',
    '0xf0e0':'◊', '0xf0e1':'⟨', '0xf0e2':'®', '0xf0e3':'©',
    '0xf0e4':'™', '0xf0e5':'∑', '0xf0e6':'⎛', '0xf0e7':'⎜',
    '0xf0e8':'⎝', '0xf0e9':'⎡', '0xf0ea':'⎢', '0xf0eb':'⎣',
    '0xf0ec':'⎧', '0xf0ed':'⎨', '0xf0ee':'⎩', '0xf0ef':'⎪',
    '0xf0f1':'⟩', '0xf0f2':'∫', '0xf0f3':'⌠',
    '0xf0f4':'⎮', '0xf0f5':'⌡', '0xf0f6':'⎞', '0xf0f7':'⎟',
    '0xf0f8':'⎠', '0xf0f9':'⎤', '0xf0fa':'⎥', '0xf0fb':'⎦',
    '0xf0fc':'⎫', '0xf0fd':'⎬', '0xf0fe':'⎭',
  }

class CharHits:
  def __init__(self):
    self._hits: Dict[str, int] = {}
  
  def add(self, char: str):
    if char in self._hits:
      self._hits[char] += 1
    else:
      self._hits[char] = 1
  
  def get_hits(self) -> List[Tuple[str, int]]:
    return sorted(self._hits.items(), key=lambda x: x[1], reverse=True)

class ConversionResult:
  def __init__(self):
    self._unconverted: CharHits = CharHits()
    self._converted: CharHits = CharHits()
    
  @property
  def unconverted(self) -> List[Tuple[str, int]]:
    return self._unconverted.get_hits()
  
  @property
  def converted(self) -> List[Tuple[str, int]]:
    return self._converted.get_hits()
    
  def add_unconverted(self, char: str):
    self._unconverted.add(char)
    
  def add_converted(self, char: str):
    self._converted.add(char)

def convert_symbols(
  doc: Document,
  font_name: str = 'Arial',
  additional_symbol_map: Dict[str, str] = {},
) -> ConversionResult:
  symbol_map = get_symbol_map()
  symbol_map.update(additional_symbol_map)
  
  result = ConversionResult()
  
  def _process_paragraph(para: Paragraph):
    previous_font_name = None
    for run in para.runs:
      if run.font.name == 'Symbol':
        is_ok = True
        text_list = list(run.text)
        for i, ch in enumerate(text_list):
          hex_code = hex(ord(ch))
          if not hex_code.startswith('0xf0'):
            continue
          if hex_code in symbol_map:
            text_list[i] = symbol_map[hex_code]
            result.add_converted(ch)
          else:
            is_ok = False
            result.add_unconverted(ch)
        run.text = ''.join(text_list)
        if is_ok:
          if previous_font_name and previous_font_name != 'Symbol':
            run.font.name = previous_font_name
          else:
            run.font.name = font_name
      else:
        previous_font_name = run.font.name
  
  for para in doc.paragraphs:
    _process_paragraph(para)
    
  for table in doc.tables:
    for row in table.rows:
      for cell in row.cells:
        for para in cell.paragraphs:
          _process_paragraph(para)
  
  return result