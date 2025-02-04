import docx
import os
import re
import cv2
import numpy as np

class ManagerDocx:
    @staticmethod
    def replace_words(path: str, old: str, new: str) -> None:
        """
        替换Word文档中的关键词。

        :param path: 文件路径
        :param old: 需要替换的关键词
        :param new: 新的替换后的关键词
        :return: None
        """
        if path.endswith(".docx"):
            # 不支持读取doc格式的文件
            doc = docx.Document(path)
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if run.text:
                        run.text = run.text.replace(old, new)
                doc.save(path)
        else:
            raise ValueError("只支持docx文件格式!")
        return path

    @staticmethod
    def change_forward(word_path: str, save_path: str) -> None:
        """
        更改Word文档的页面方向。

        :param word_path: Word文件路径
        :param result_path: 结果文件路径
        :return: None
        """
        doc = docx.Document(word_path)
        for section in doc.sections:
            # 交替宽高
            section.page_width, section.page_height = section.page_height, section.page_width
        # 保存为新文件
        doc.save(save_path) 
        return save_path

    @staticmethod
    def get_pictures(word_path: str, result_path: str) -> str:
        """
        从Word文档中提取图片并保存。

        :param word_path: Word文件路径
        :param result_path: 图片保存路径
        :return: 图片保存路径
        """
        # 创建保存路径
        if not os.path.exists(result_path):
            os.makedirs(result_path)
        # 读取文件
        doc = docx.Document(word_path)

        # 获取图片
        dict_rel = doc.part._rels
        for rel in dict_rel:
            rel = dict_rel[rel]
            if "image" in rel.target_ref:            
                img_name = re.findall("/(.*)", rel.target_ref)[0]
                word_name = os.path.splitext(word_path)[0]
                if os.sep in word_name:
                    new_name = word_name.split('\\')[-1]
                else:
                    new_name = word_name.split('/')[-1]
                # cv2获取图片大小
                imgdata = np.frombuffer(rel.target_part.blob,np.uint8)
                img_cv = cv2.imdecode(imgdata,cv2.IMREAD_ANYCOLOR)
                img_name = '{}-{}-{}-{}'.format(new_name,img_cv.shape[0],img_cv.shape[1],img_name)
                # 直接二进制写入兼容性比使用CV2的保存图片好
                with open(f'{result_path}/{img_name}','wb') as f:
                    f.write(rel.target_part.blob)
            else:
                pass
        return result_path
    
